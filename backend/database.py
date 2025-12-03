import os
import time
import uuid
from sqlite3 import connect
from docx import Document

DATABASE_PATH = "database/database.db"


def get_db_connection(db_path=DATABASE_PATH):
    # 确保目录存在（避免连接时路径不存在）
    dirpath = os.path.dirname(db_path)
    if dirpath and not os.path.exists(dirpath):
        os.makedirs(dirpath, exist_ok=True)
    conn = connect(db_path)
    print("Database connection established.", db_path)
    return conn


def close_db_connection(conn):
    conn.close()
    print("Database connection closed.")


def init_db(db_path=DATABASE_PATH):
    connection = get_db_connection(db_path)
    cursor = connection.cursor()
    # 使用 IF NOT EXISTS 避免重复创建导致错误
    cursor.execute(
        "CREATE TABLE IF NOT EXISTS chat_history(person TEXT, message TEXT, talkcycle INTEGER)"
    )
    connection.commit()
    close_db_connection(connection)


def update_db(role: str, message: str, talkcycle: int, db_path=DATABASE_PATH):
    connection = get_db_connection(db_path)
    cursor = connection.cursor()
    cursor.execute(
        "INSERT INTO chat_history (person, message, talkcycle) VALUES (?, ?, ?)",
        (role, message, talkcycle),
    )
    connection.commit()
    close_db_connection(connection)


def get_all_chat_history(db_path=DATABASE_PATH):
    connection = get_db_connection(db_path)
    cursor = connection.cursor()
    cursor.execute("SELECT * FROM chat_history")
    rows = cursor.fetchall()
    close_db_connection(connection)
    return rows


def transform_db_to_word(db_path=DATABASE_PATH):
    """
    导出数据库聊天记录为 .docx 并返回文件信息字典：
    { "file_path": "...", "filename": "chat_history_....docx" }
    """
    rows = get_all_chat_history(db_path)
    doc = Document()
    doc.add_heading("聊天记录导出", level=1)

    if not rows:
        doc.add_paragraph("No chat history found.")

    for row in rows:
        person, message, talkcycle = row
        doc.add_paragraph(f"轮次 {talkcycle} - {person}:")
        for line in str(message).splitlines():
            doc.add_paragraph(line)

    # 输出到数据库目录下的 temp 子目录（确保路径存在）
    base_dir = os.path.dirname(db_path) or "."
    out_dir = os.path.join(base_dir, "exports")
    os.makedirs(out_dir, exist_ok=True)

    # 唯一文件名，避免缓存/并发冲突
    filename = f"chat_history_{time.strftime('%Y%m%d_%H%M%S', time.localtime())}.docx"
    out_path = os.path.abspath(os.path.join(out_dir, filename))

    doc.save(out_path)
    print(f"transform_db_to_word: saved {out_path} ({len(rows)} rows)")

    return {"file_path": out_path, "filename": filename}


init_db()  # 确保首次导入时初始化数据库结构
